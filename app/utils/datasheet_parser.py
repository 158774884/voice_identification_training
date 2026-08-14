"""
Datasheet parser — extract chip specs from a datasheet file (PDF / txt / docx).

Offline, rule-based: pulls text out of the document and applies regex/keyword
heuristics to fill a :class:`ChipSpec` for evaluation. Results are meant to be
reviewed by the user in the chip editor dialog before saving.

解析策略:
- 关键词上下文锚定: 只在 CPU/核心、SRAM、Flash、功耗等关键词附近取值
- 排除干扰: 2.4/5GHz 无线频段、PSRAM/DRAM、TX/RF 发射功率等
- 取典型值: 功耗优先 "typical/工作/运行" 上下文，缺失时用 "mA x V" 估算
"""
import os
import re
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

from app.models.chip_database import ChipSpec


# Fields we consider "important" when scoring how much was extracted.
IMPORTANT_FIELDS = [
    "manufacturer",
    "cpu_freq_mhz",
    "cpu_cores",
    "ram_kb",
    "flash_kb",
    "npu_tops",
    "power_consumption_mw",
]

_UNIT_TO_KB = {"kb": 1.0, "mb": 1024.0, "gb": 1024.0 * 1024.0}

_MANUFACTURERS = {
    "STMicroelectronics": ["stmicroelectronics", "st microelectronics"],
    "NXP": ["nxp semiconductors", "nxp"],
    "Texas Instruments": ["texas instruments"],
    "Espressif": ["espressif", "乐鑫"],
    "Qualcomm": ["qualcomm"],
    "Rockchip": ["rockchip", "瑞芯微"],
    "Allwinner": ["allwinner", "全志"],
    "Nordic": ["nordic semiconductor", "nordic"],
    "Realtek": ["realtek", "瑞昱"],
    "Ambiq": ["ambiq"],
    "Infineon": ["infineon"],
    "Renesas": ["renesas"],
    "GigaDevice": ["gigadevice", "兆易创新"],
    "Synaptics": ["synaptics"],
    "MediaTek": ["mediatek", "联发科"],
    "HiSilicon": ["hisilicon", "海思"],
    "JieLi": ["jieli", "杰理", "珠海杰理"],
    "Amlogic": ["amlogic", "晶晨"],
    "Beken": ["beken", "博通集成"],
    "Kendryte": ["kendryte", "嘉楠"],
    "Bouffalo": ["bouffalo", "博流"],
    "XMOS": ["xmos"],
    "NVIDIA": ["nvidia"],
    "Axera": ["axera", "爱芯元智"],
}

_CORE_WORDS_EN = {
    "single": 1, "dual": 2, "triple": 3, "quad": 4,
    "hexa": 6, "hex": 6, "octa": 8, "octo": 8, "deca": 10,
}
_CORE_WORDS_CN = {"单核": 1, "双核": 2, "四核": 4, "六核": 6, "八核": 8}

# 无线射频频段 (蓝牙/WiFi)，不是 CPU 主频
_RADIO_BAND_GHZ = {2.4, 2.5, 5.0, 5.8, 24.0, 60.0}

_CPU_CTX = re.compile(r'cpu|risc|core|cortex|arm|mcu|主频|处理器|时钟|核|clock', re.IGNORECASE)


@dataclass
class DatasheetParseResult:
    """Outcome of parsing a datasheet file."""

    chip: ChipSpec = field(default_factory=ChipSpec)
    confidence: float = 0.0          # 0.0 - 1.0, fraction of important fields found
    found_fields: List[str] = field(default_factory=list)
    excerpt: str = ""                # short snippet of matched source text


# ================================================================
# Text extraction
# ================================================================

def extract_text(path: str) -> str:
    """Extract plain text from a datasheet file, dispatching by extension."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    return _extract_plaintext(path)


def _extract_pdf(path: str) -> str:
    errors = []
    try:
        import fitz
        doc = fitz.open(path)
        try:
            return "\n".join(page.get_text() for page in doc)
        finally:
            doc.close()
    except Exception as e:  # noqa: BLE001
        errors.append(f"fitz: {e}")

    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join((page.extract_text() or "") for page in pdf.pages)
    except Exception as e:  # noqa: BLE001
        errors.append(f"pdfplumber: {e}")

    try:
        from pypdf import PdfReader
        return "\n".join((page.extract_text() or "") for page in PdfReader(path).pages)
    except Exception as e:  # noqa: BLE001
        errors.append(f"pypdf: {e}")

    raise RuntimeError("无法解析 PDF（缺少 PDF 库或文件损坏）: " + "; ".join(errors))


def _extract_docx(path: str) -> str:
    import docx
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_plaintext(path: str) -> str:
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


# ================================================================
# Field extractors (each returns (value_or_None, snippet_or_empty))
# ================================================================

def _to_kb(value: float, unit: str) -> float:
    return value * _UNIT_TO_KB.get(unit.lower(), 1.0)


def _median(values: List[float]) -> float:
    """Lower median — avoids outliers that a naive max() would pick."""
    s = sorted(values)
    return s[(len(s) - 1) // 2]


def _extract_freq(text: str) -> Tuple[Optional[int], str]:
    """CPU 主频: 排除无线频段，优先 CPU/核心上下文。"""
    candidates = []  # (mhz, snippet, is_cpu_ctx)
    for m in re.finditer(r'(\d+(?:\.\d+)?)\s*(MHz|GHz)', text, re.IGNORECASE):
        v = float(m.group(1))
        unit = m.group(2).lower()
        mhz = v * 1000.0 if unit == 'ghz' else v
        # 2.4/5/24 GHz 是蓝牙/WiFi 频段，不是 CPU 主频
        if unit == 'ghz' and round(v, 1) in _RADIO_BAND_GHZ:
            continue
        s, e = m.start(), m.end()
        window = text[max(0, s - 24): min(len(text), e + 24)]
        candidates.append((mhz, m.group(0), bool(_CPU_CTX.search(window))))

    if not candidates:
        return None, ""

    cpu_cands = [c for c in candidates if c[2]]
    pool = cpu_cands or candidates
    target = _median([c[0] for c in pool])
    for c in pool:
        if c[0] == target:
            return int(round(target)), c[1]
    return int(round(target)), pool[0][1]


def _extract_cores(text: str) -> Optional[int]:
    t = text.lower()
    for w, n in _CORE_WORDS_CN.items():
        if w in t:
            return n
    for w, n in _CORE_WORDS_EN.items():
        if re.search(rf'\b{w}\s*-?\s*core', t):
            return n
    m = re.search(r'(\d+)\s*-?\s*core', t)
    if m:
        return int(m.group(1))
    return None


def _extract_memory(text: str, keywords: List[str]) -> Tuple[Optional[int], str]:
    """Find a memory size (KB/MB/GB) near any `keywords` (either order).
    Word-bounded so "PSRAM"/"DRAM"/"SDRAM" don't match "SRAM"/"RAM".
    """
    kw = "|".join(keywords)
    results = []
    pat_a = re.compile(r'(\d+(?:\.\d+)?)\s*(KB|MB|GB)\b[\s:：]*\b(' + kw + r')\b', re.IGNORECASE)
    for m in pat_a.finditer(text):
        results.append((_to_kb(float(m.group(1)), m.group(2)), m.group(0)))
    pat_b = re.compile(r'\b(' + kw + r')\b[\s:：]*(\d+(?:\.\d+)?)\s*(KB|MB|GB)\b', re.IGNORECASE)
    for m in pat_b.finditer(text):
        results.append((_to_kb(float(m.group(2)), m.group(3)), m.group(0)))
    if not results:
        return None, ""
    target = _median([r[0] for r in results])
    for r in results:
        if r[0] == target:
            return int(round(target)), r[1]
    return int(round(target)), results[0][1]


def _extract_npu(text: str) -> Tuple[Optional[float], str]:
    tops = []
    for m in re.finditer(r'(\d+(?:\.\d+)?)\s*TOPS', text, re.IGNORECASE):
        tops.append((float(m.group(1)), m.group(0)))
    for m in re.finditer(r'(\d+(?:\.\d+)?)\s*GOPS', text, re.IGNORECASE):
        tops.append((float(m.group(1)) / 1000.0, m.group(0)))
    if not tops:
        return None, ""
    best = max(tops, key=lambda t: t[0])
    return best[0], best[1]


def _extract_power_from_current(text: str) -> Tuple[Optional[int], str]:
    """Fallback: 工作电流 (mA) x 供电电压 (V) 估算功耗 (mW)."""
    ma_matches = list(re.finditer(r'(\d+(?:\.\d+)?)\s*mA\b', text, re.IGNORECASE))
    if not ma_matches:
        return None, ""
    # 排除待机/休眠的小电流 (<1mA)
    ma_vals = [(float(m.group(1)), m) for m in ma_matches]
    active = [x for x in ma_vals if x[0] >= 1.0] or ma_vals
    ma, m = max(active, key=lambda x: x[0])
    v_match = re.search(r'(\d+(?:\.\d+)?)\s*V\b', text)
    voltage = float(v_match.group(1)) if v_match else 3.3
    return int(round(ma * voltage)), m.group(0)


def _extract_power(text: str) -> Tuple[Optional[int], str]:
    """功耗: 优先 typical/工作/运行 上下文，排除 TX/RF 发射功率。"""
    candidates = []
    for m in re.finditer(r'(\d+(?:\.\d+)?)\s*(mW|W|uW|µW|μW)\b', text, re.IGNORECASE):
        v = float(m.group(1))
        unit = m.group(2).lower()
        if unit == "mw":
            mw = v
        elif unit == "w":
            mw = v * 1000.0
        else:
            mw = v / 1000.0
        s, e = m.start(), m.end()
        window = text[max(0, s - 30): min(len(text), e + 30)]
        if re.search(r'\btx\b|\brf\b|transmit|发射|射频|\bpa\b|功放|amplifier', window, re.IGNORECASE):
            continue
        is_typical = bool(re.search(r'typical|active|normal|power|功耗|功率|典型|工作|运行|电流',
                                    window, re.IGNORECASE))
        candidates.append((mw, m.group(0), is_typical))

    if not candidates:
        return _extract_power_from_current(text)

    pool = [c for c in candidates if c[2]] or candidates
    target = _median([c[0] for c in pool])
    for c in pool:
        if c[0] == target:
            return int(round(target)), c[1]
    return int(round(target)), pool[0][1]


def _extract_manufacturer(text: str) -> str:
    t = text.lower()
    for mfr, keys in _MANUFACTURERS.items():
        if any(k in t for k in keys):
            return mfr
    return ""


def _extract_architecture(text: str) -> str:
    t = text.lower()
    if re.search(r'cortex-?\s*a\d', t):
        return "SoC"
    if re.search(r'cortex-?\s*m\d|\brisc-?v\b', t):
        return "MCU"
    if re.search(r'\bnpu\b|\btpu\b|\bneural', t):
        return "NPU"
    if re.search(r'\bdsp\b|\bhifi\b|\bcadence\b', t):
        return "DSP"
    return "MCU"


def _has_dsp(text: str) -> bool:
    return bool(re.search(r'\bdsp\b|\bhifi\b|\bcadence\b', text, re.IGNORECASE))


def _has_mva(text: str) -> bool:
    return bool(re.search(r'\bmva\b|矩阵|向量加速', text, re.IGNORECASE))


def _extract_name(path: str) -> str:
    stem = os.path.splitext(os.path.basename(path))[0]
    stem = re.sub(r'[_\-]+', ' ', stem)
    stem = re.sub(r'\b(datasheet|data[- ]?sheet|manual|ds|规格书|数据手册)\b', '', stem,
                  flags=re.IGNORECASE)
    stem = re.sub(r'\s+', ' ', stem).strip(' -_')
    return stem or "UnknownChip"


# ================================================================
# Public API
# ================================================================

def parse_text(text: str, filename: str = "") -> DatasheetParseResult:
    """Parse plain datasheet text into a DatasheetParseResult."""
    found: dict = {}
    snippets: List[str] = []

    name = _extract_name(filename) if filename else ""
    mfr = _extract_manufacturer(text)
    if mfr:
        found["manufacturer"] = mfr

    freq, freq_snip = _extract_freq(text)
    if freq is not None:
        found["cpu_freq_mhz"] = freq
        snippets.append(freq_snip)

    cores = _extract_cores(text)
    if cores is not None:
        found["cpu_cores"] = cores

    ram, ram_snip = _extract_memory(text, ["SRAM", "RAM"])
    if ram is not None:
        found["ram_kb"] = ram
        snippets.append(ram_snip)

    flash, flash_snip = _extract_memory(text, ["Flash", "FLASH"])
    if flash is not None:
        found["flash_kb"] = flash
        snippets.append(flash_snip)

    npu, npu_snip = _extract_npu(text)
    if npu is not None:
        found["npu_tops"] = npu
        snippets.append(npu_snip)

    if _has_dsp(text):
        found["dsp"] = True

    power, power_snip = _extract_power(text)
    if power is not None:
        found["power_consumption_mw"] = power
        snippets.append(power_snip)

    arch = _extract_architecture(text)

    notes = f"从数据手册解析: {filename}" if filename else "从数据手册解析"
    if _has_mva(text):
        notes += "；检测到 MVA/矩阵向量加速器（非 TOPS 级 NPU）"

    chip = ChipSpec(
        name=name,
        manufacturer=mfr,
        architecture=arch,
        cpu_cores=found.get("cpu_cores", 1),
        cpu_freq_mhz=found.get("cpu_freq_mhz", 100),
        ram_kb=found.get("ram_kb", 128),
        flash_kb=found.get("flash_kb", 1024),
        npu_tops=found.get("npu_tops", 0.0),
        dsp=found.get("dsp", False),
        power_consumption_mw=found.get("power_consumption_mw", 500),
        max_model_size_kb=found.get("flash_kb", 1024),
        notes=notes,
    )

    found_fields = [f for f in IMPORTANT_FIELDS if f in found]
    confidence = len(found_fields) / len(IMPORTANT_FIELDS)

    return DatasheetParseResult(
        chip=chip,
        confidence=confidence,
        found_fields=found_fields,
        excerpt="\n".join(dict.fromkeys(snippets)),
    )


def parse_datasheet(path: str) -> DatasheetParseResult:
    """Extract text from a datasheet file and parse chip specs from it."""
    text = extract_text(path)
    return parse_text(text, filename=os.path.basename(path))
